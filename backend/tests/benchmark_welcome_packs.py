"""
Welcome Pack benchmark harness — tests the full pipeline for all 5 sample leases,
downloads each generated Welcome Pack, runs automated checks, and saves the
.docx files for manual review.

Usage:
    python tests/benchmark_welcome_packs.py

Requires:
    - Backend running at http://localhost:8000
    - Test user credentials in environment or .env
    - All 5 sample lease files in ../template/
"""

import json
import os
import sys
import time
from pathlib import Path

import httpx
from docx import Document

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

API_URL = os.getenv("API_URL", "http://localhost:8000")
SCRIPT_DIR = Path(__file__).parent
TEMPLATE_DIR = SCRIPT_DIR.parent.parent / "template"
TEST_RESULTS_DIR = SCRIPT_DIR.parent.parent / "docs" / "test_results"
GROUND_TRUTH_PATH = SCRIPT_DIR / "ground_truth.json"

# Leases where special_conditions is null — section must be removed
NULL_SPECIAL_CONDITIONS = {
    "Lease Agreement - David Okafor.docx",
    "Lease Agreement - Sarah Chen.docx",
    "Lease Agreement - Marcus & Lisa Johnson.docx",
}

# Leases where special_conditions has content — section must be present
HAS_SPECIAL_CONDITIONS = {
    "Lease Agreement - Emma Whitfield.docx",
    "Lease Agreement - Raj Patel.docx",
}

LEASE_FILES = [
    "Lease Agreement - David Okafor.docx",
    "Lease Agreement - Sarah Chen.docx",
    "Lease Agreement - Emma Whitfield.docx",
    "Lease Agreement - Raj Patel.docx",
    "Lease Agreement - Marcus & Lisa Johnson.docx",
]


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def load_env():
    """Load env vars from backend/.env if not already set."""
    env_path = SCRIPT_DIR.parent / ".env"
    if env_path.exists():
        for line in env_path.read_text().splitlines():
            line = line.strip()
            if line and not line.startswith("#") and "=" in line:
                key, _, value = line.partition("=")
                os.environ.setdefault(key.strip(), value.strip())

    # Try frontend .env for anon key
    frontend_env = SCRIPT_DIR.parent.parent / "frontend" / ".env"
    if frontend_env.exists():
        for line in frontend_env.read_text().splitlines():
            if line.startswith("VITE_SUPABASE_ANON_KEY="):
                os.environ.setdefault("SUPABASE_ANON_KEY", line.split("=", 1)[1].strip())
                break


def get_jwt_token() -> str:
    """Authenticate against Supabase and return an access token."""
    supabase_url = os.getenv("SUPABASE_URL", "")
    anon_key = os.getenv("SUPABASE_ANON_KEY", "")
    resp = httpx.post(
        f"{supabase_url}/auth/v1/token?grant_type=password",
        json={"email": os.getenv("TEST_EMAIL", "test@acmepg.com.au"),
              "password": os.getenv("TEST_PASSWORD", "TestPassword123!")},
        headers={"Content-Type": "application/json", "apikey": anon_key},
        timeout=15,
    )
    resp.raise_for_status()
    return resp.json()["access_token"]


def get_all_text(doc: Document) -> str:
    """Get all text from paragraphs and tables."""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return " ".join(parts)


def count_occurrences(doc: Document, text: str) -> int:
    """Count how many times text appears across all paragraphs and table cells."""
    count = 0
    for p in doc.paragraphs:
        count += p.text.count(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                count += cell.text.count(text)
    return count


# ---------------------------------------------------------------------------
# Checks
# ---------------------------------------------------------------------------

def check_no_leftover_placeholders(doc: Document) -> tuple[bool, str]:
    """Verify no {{...}} placeholders remain."""
    leftover = []
    for i, p in enumerate(doc.paragraphs):
        if "{{" in p.text:
            leftover.append(f"P[{i}]: {p.text[:60]}")
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if "{{" in cell.text:
                    leftover.append(f"T{t_idx}[{r_idx},{c_idx}]: {cell.text[:60]}")

    if leftover:
        return False, f"Found {len(leftover)} leftover: {'; '.join(leftover[:3])}"
    return True, "No leftover placeholders"


def check_tenant_name_occurrences(doc: Document, expected_name: str) -> tuple[bool, str]:
    """Verify tenant_name appears in all 3 expected locations."""
    count = count_occurrences(doc, expected_name)
    if count >= 3:
        return True, f"'{expected_name}' found {count} times (expected >= 3)"
    return False, f"'{expected_name}' found only {count} times (expected >= 3)"


def check_property_address_occurrences(doc: Document, expected_address: str) -> tuple[bool, str]:
    """Verify property_address appears in all 4 expected locations."""
    # Use a shorter substring for matching — the full address may wrap differently
    short_addr = expected_address.split(",")[0]  # e.g. "42 Oakwood Avenue"
    count = count_occurrences(doc, short_addr)
    if count >= 4:
        return True, f"'{short_addr}' found {count} times (expected >= 4)"
    # If short match works but full doesn't, still report
    full_count = count_occurrences(doc, expected_address)
    if full_count >= 4:
        return True, f"Full address found {full_count} times (expected >= 4)"
    return False, f"'{short_addr}' found {count} times, full address {full_count} times (expected >= 4)"


def check_special_conditions_removed(doc: Document) -> tuple[bool, str]:
    """Verify Special Conditions heading and placeholder are physically removed."""
    for p in doc.paragraphs:
        text = p.text.strip()
        if text == "Special Conditions":
            return False, "Heading 'Special Conditions' still present"
        if "{{special_conditions}}" in text:
            return False, "Placeholder {{special_conditions}} still present"
    return True, "Section physically removed from XML"


def check_special_conditions_present(doc: Document) -> tuple[bool, str]:
    """Verify Special Conditions section exists with real content."""
    found_heading = False
    found_content = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if text == "Special Conditions":
            found_heading = True
        if found_heading and text and text != "Special Conditions" and "{{" not in text:
            found_content = True
            break

    if found_heading and found_content:
        return True, "Section present with content"
    if not found_heading:
        return False, "Heading 'Special Conditions' not found"
    return False, "Heading found but no content after it"


def check_value_present(doc: Document, label: str, value: str) -> tuple[bool, str]:
    """Check that a specific value appears in the document."""
    all_text = get_all_text(doc)
    if value in all_text:
        return True, f"'{value}' found"
    return False, f"'{value}' NOT found"


# ---------------------------------------------------------------------------
# Main benchmark
# ---------------------------------------------------------------------------

def run_benchmark():
    load_env()

    # Load ground truth for expected values
    with open(GROUND_TRUTH_PATH) as f:
        ground_truth = json.load(f)

    TEST_RESULTS_DIR.mkdir(parents=True, exist_ok=True)

    print("=" * 70)
    print("WELCOME PACK BENCHMARK — Full Pipeline Test")
    print(f"API: {API_URL}")
    print(f"Leases: {len(LEASE_FILES)}")
    print(f"Output: {TEST_RESULTS_DIR}")
    print("=" * 70)

    # Authenticate
    print("\nAuthenticating...", end=" ")
    token = get_jwt_token()
    print("OK")

    total_pass = 0
    total_fail = 0
    total_checks = 0
    failures: list[str] = []
    latencies: list[float] = []

    for file_name in LEASE_FILES:
        file_path = TEMPLATE_DIR / file_name
        if not file_path.exists():
            print(f"\nSKIPPED: {file_name} — file not found")
            continue

        expected = ground_truth.get(file_name, {})
        tenant_name = expected.get("tenant_name", "Unknown")
        property_address = expected.get("property_address", "Unknown")

        print(f"\n{'─' * 70}")
        print(f"LEASE: {file_name}")
        print(f"{'─' * 70}")

        # Upload and process
        start = time.time()
        with open(file_path, "rb") as f:
            resp = httpx.post(
                f"{API_URL}/api/lease/upload",
                headers={"Authorization": f"Bearer {token}"},
                files={"file": (file_name, f, "application/octet-stream")},
                timeout=120,
            )
        latency = time.time() - start
        latencies.append(latency)

        if resp.status_code != 200:
            print(f"  ERROR: HTTP {resp.status_code} — {resp.text[:200]}")
            failures.append(f"{file_name}: HTTP {resp.status_code}")
            continue

        result = resp.json()
        upload_id = result.get("upload_id")
        pipeline_status = result.get("status")
        print(f"  Pipeline: {pipeline_status} | Latency: {latency:.1f}s")

        if pipeline_status != "complete":
            print(f"  ERROR: Expected status 'complete', got '{pipeline_status}'")
            failures.append(f"{file_name}: status={pipeline_status}")
            continue

        # Download Welcome Pack
        dl_resp = httpx.get(
            f"{API_URL}/api/welcome-pack/download/{upload_id}",
            headers={"Authorization": f"Bearer {token}"},
            timeout=30,
        )

        if dl_resp.status_code != 200:
            print(f"  ERROR: Download failed — HTTP {dl_resp.status_code}")
            failures.append(f"{file_name}: download HTTP {dl_resp.status_code}")
            continue

        # Save .docx for manual review
        safe_name = f"Welcome_Pack_{tenant_name.replace(' ', '_').replace('&', 'and')}.docx"
        save_path = TEST_RESULTS_DIR / safe_name
        save_path.write_bytes(dl_resp.content)
        print(f"  Saved: {safe_name} ({len(dl_resp.content)} bytes)")

        # Load document for automated checks
        doc = Document(str(save_path))
        print()

        # Run checks
        checks: list[tuple[str, tuple[bool, str]]] = []

        # 1. No leftover placeholders
        checks.append(("No leftover {{...}}", check_no_leftover_placeholders(doc)))

        # 2. Tenant name occurrences (>= 3)
        checks.append(("tenant_name x3", check_tenant_name_occurrences(doc, tenant_name)))

        # 3. Property address occurrences (>= 4)
        checks.append(("property_address x4", check_property_address_occurrences(doc, property_address)))

        # 4. Special conditions
        if file_name in NULL_SPECIAL_CONDITIONS:
            checks.append(("special_conditions removed", check_special_conditions_removed(doc)))
        elif file_name in HAS_SPECIAL_CONDITIONS:
            checks.append(("special_conditions present", check_special_conditions_present(doc)))

        # 5. Lease-specific checks
        if file_name == "Lease Agreement - Marcus & Lisa Johnson.docx":
            checks.append(("Joint tenancy names", check_value_present(doc, "joint tenancy", "Marcus Johnson & Lisa Johnson")))

        if file_name == "Lease Agreement - Raj Patel.docx":
            checks.append(("Fortnightly rent", check_value_present(doc, "fortnightly rent", "$1,150.00 per fortnight")))
            checks.append(("Parking details", check_value_present(doc, "parking space", "Space #47")))

        if file_name == "Lease Agreement - David Okafor.docx":
            checks.append(("Pet conditions detail", check_value_present(doc, "pet conditions", "desexed")))

        # Key contacts
        checks.append(("Property manager name", check_value_present(doc, "PM name", "Julia Torres")))
        checks.append(("Property manager email", check_value_present(doc, "PM email", "julia.torres@acmepg.com.au")))

        # Report
        lease_pass = 0
        lease_fail = 0
        for label, (passed, detail) in checks:
            total_checks += 1
            if passed:
                lease_pass += 1
                total_pass += 1
                print(f"  ✓ {label}: {detail}")
            else:
                lease_fail += 1
                total_fail += 1
                print(f"  ✗ {label}: {detail}")
                failures.append(f"{file_name} → {label}: {detail}")

        print(f"\n  Result: {lease_pass}/{lease_pass + lease_fail} checks passed")

    # Summary
    print(f"\n{'=' * 70}")
    print("SUMMARY")
    print(f"{'=' * 70}")
    print(f"  Total: {total_pass}/{total_checks} checks passed")
    print(f"  Avg latency: {sum(latencies) / len(latencies):.1f}s" if latencies else "  No latency data")
    print(f"  Min latency: {min(latencies):.1f}s" if latencies else "")
    print(f"  Max latency: {max(latencies):.1f}s" if latencies else "")

    print(f"\n  Generated Welcome Packs saved to: {TEST_RESULTS_DIR}")

    if failures:
        print(f"\n  FAILURES ({len(failures)}):")
        for fail in failures:
            print(f"    ✗ {fail}")

    print()
    if total_fail == 0:
        print(f"  ✓ ALL CHECKS PASSED — {total_pass}/{total_checks}")
        print("  Note: Open the .docx files to visually verify formatting and branding")
        return 0
    else:
        print(f"  ✗ {total_fail} CHECK(S) FAILED")
        return 1


if __name__ == "__main__":
    sys.exit(run_benchmark())
