"""
Welcome Pack .docx generation service.

Loads the Tenant Welcome Pack Template, replaces all 14 {{placeholder}} markers
with extracted lease data, and conditionally removes the Special Conditions
section when special_conditions is null.

Handles python-docx run-splitting: placeholders may be split across multiple
XML runs within a paragraph. The replacement logic works at the paragraph level
and preserves formatting from the original placeholder runs.
"""

import io
import logging
import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from app.config import settings

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Template path — resolved from settings (env var TEMPLATE_PATH or local default)
# ---------------------------------------------------------------------------

TEMPLATE_PATH = Path(settings.template_path)

# ---------------------------------------------------------------------------
# Field name → template placeholder mapping
# ---------------------------------------------------------------------------

PLACEHOLDER_MAP = {
    "tenant_name": "{{tenant_name}}",
    "property_address": "{{property_address}}",
    "lease_start_date": "{{lease_start_date}}",
    "lease_end_date": "{{lease_end_date}}",
    "rent_amount": "{{rent_amount}}",
    "bond_amount": "{{bond_amount}}",
    "num_occupants": "{{num_occupants}}",
    "pet_permission": "{{pet_permission}}",
    "parking": "{{parking_included}}",  # template uses different name
    "special_conditions": "{{special_conditions}}",
    "landlord_name": "{{landlord_name}}",
    "property_manager_name": "{{property_manager_name}}",
    "property_manager_email": "{{property_manager_email}}",
    "property_manager_phone": "{{property_manager_phone}}",
}

# Paragraph indices for the Special Conditions section
_SPECIAL_CONDITIONS_HEADING_IDX = 11  # "Special Conditions" heading
_SPECIAL_CONDITIONS_BODY_IDX = 12     # {{special_conditions}} placeholder


# ---------------------------------------------------------------------------
# Rent normalization
# ---------------------------------------------------------------------------

def _normalize_rent(rent_str: str) -> str:
    """
    Normalize rent to consistent monthly format for the Welcome Pack.

    Converts fortnightly (x 2.17262) and weekly (x 4.35) amounts to monthly.
    Strips 'AUD' and 'calendar' if Gemini still includes them.
    Monthly amounts pass through unchanged.

    Returns format: '$X,XXX.XX per month'
    """
    if not rent_str:
        return rent_str

    # Strip AUD and "calendar" if Gemini still includes them
    cleaned = rent_str.replace(" AUD", "").replace("AUD ", "").replace("calendar ", "").strip()

    # Extract dollar amount and frequency
    match = re.match(r"\$?([\d,]+\.?\d*)\s+per\s+(\w+)", cleaned)
    if not match:
        logger.warning("Could not parse rent_amount '%s' — returning as-is", rent_str)
        return rent_str

    amount_str = match.group(1).replace(",", "")
    frequency = match.group(2).lower()

    try:
        amount = float(amount_str)
    except ValueError:
        return rent_str

    # Convert to monthly
    if frequency == "fortnight":
        amount = amount * 2.17262
    elif frequency == "week":
        amount = amount * 4.35
    elif frequency == "month":
        pass  # Already monthly
    else:
        logger.warning("Unknown rent frequency '%s' — returning as-is", frequency)
        return rent_str

    normalized = f"${amount:,.2f} per month"
    if frequency != "month":
        logger.info("Converted rent from '%s' to '%s'", rent_str, normalized)
    return normalized


# ---------------------------------------------------------------------------
# Run-splitting-safe paragraph replacement
# ---------------------------------------------------------------------------

def _replace_in_paragraph(paragraph, placeholder: str, replacement: str) -> bool:
    """
    Replace a placeholder in a paragraph, handling run-splitting.

    If the placeholder is split across multiple runs, this function:
    1. Finds which runs contain parts of the placeholder
    2. Puts the full replacement text in the first matching run
    3. Clears the remaining runs that were part of the placeholder
    4. Preserves the formatting of the first matching run

    Returns True if a replacement was made.
    """
    # Fast check: is the placeholder even in this paragraph?
    full_text = "".join(run.text for run in paragraph.runs)
    if placeholder not in full_text:
        return False

    # Build a map of character positions to runs
    runs = paragraph.runs
    if not runs:
        return False

    # Find all occurrences and replace them
    replaced = False
    while True:
        # Recalculate full text each iteration (runs may have changed)
        full_text = "".join(run.text for run in runs)
        start_pos = full_text.find(placeholder)
        if start_pos == -1:
            break

        end_pos = start_pos + len(placeholder)
        replaced = True

        # Map character positions to run indices
        char_offset = 0
        run_ranges = []  # [(run_idx, run_start_char, run_end_char), ...]
        for i, run in enumerate(runs):
            run_start = char_offset
            run_end = char_offset + len(run.text)
            run_ranges.append((i, run_start, run_end))
            char_offset = run_end

        # Find which runs are touched by the placeholder
        affected_runs = []
        for run_idx, run_start, run_end in run_ranges:
            if run_end > start_pos and run_start < end_pos:
                affected_runs.append((run_idx, run_start, run_end))

        if not affected_runs:
            break

        # First affected run: replace the placeholder portion with the replacement text
        first_idx, first_start, first_end = affected_runs[0]
        first_run = runs[first_idx]

        # Calculate what part of this run's text is before/after the placeholder
        cut_start = max(0, start_pos - first_start)
        cut_end = min(first_end - first_start, end_pos - first_start)

        before = first_run.text[:cut_start]
        after_in_first = first_run.text[cut_end:]

        first_run.text = before + replacement + after_in_first

        # Remaining affected runs: trim or clear the placeholder portion
        for run_idx, run_start, run_end in affected_runs[1:]:
            run = runs[run_idx]
            # How much of this run is part of the placeholder?
            trim_start = max(0, start_pos - run_start)
            trim_end = min(len(run.text), end_pos - run_start)

            run.text = run.text[:trim_start] + run.text[trim_end:]

    return replaced


def _replace_in_table(table, placeholder: str, replacement: str) -> bool:
    """Replace a placeholder across all cells in a table."""
    replaced = False
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if _replace_in_paragraph(paragraph, placeholder, replacement):
                    replaced = True
    return replaced


# ---------------------------------------------------------------------------
# Special conditions section removal
# ---------------------------------------------------------------------------

def _remove_special_conditions_section(doc: Document) -> None:
    """
    Remove the Special Conditions heading (P[11]) and body (P[12])
    from the document XML. Reverse order to prevent index shift.
    """
    for idx in [_SPECIAL_CONDITIONS_BODY_IDX, _SPECIAL_CONDITIONS_HEADING_IDX]:
        para = doc.paragraphs[idx]
        parent = para._element.getparent()
        parent.remove(para._element)
        logger.info("Removed paragraph P[%d] from document XML", idx)


# ---------------------------------------------------------------------------
# Main generation function
# ---------------------------------------------------------------------------

def generate_welcome_pack(extracted_data: dict) -> bytes:
    """
    Generate a Welcome Pack .docx from extracted lease data.

    Args:
        extracted_data: Dict with all 14 extracted fields.

    Returns:
        The generated .docx file as bytes.
    """
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Welcome Pack template not found at {TEMPLATE_PATH}")

    doc = Document(str(TEMPLATE_PATH))

    # Step 1: Handle special_conditions — remove section or keep it
    special_conditions = extracted_data.get("special_conditions")
    if special_conditions is None:
        logger.info("special_conditions is null — removing section from document")
        _remove_special_conditions_section(doc)
    else:
        logger.info("special_conditions has content — section will be populated")

    # Step 2: Build replacement dict from PLACEHOLDER_MAP
    replacements = {}
    for field_name, placeholder in PLACEHOLDER_MAP.items():
        value = extracted_data.get(field_name)
        if value is not None:
            # Normalize rent to monthly for the Welcome Pack display
            if field_name == "rent_amount":
                value = _normalize_rent(str(value))
            replacements[placeholder] = str(value)

    # Step 3: Replace placeholders in all paragraphs
    for placeholder, value in replacements.items():
        for paragraph in doc.paragraphs:
            _replace_in_paragraph(paragraph, placeholder, value)

    # Step 4: Replace placeholders in all tables
    for placeholder, value in replacements.items():
        for table in doc.tables:
            _replace_in_table(table, placeholder, value)

    # Step 5: Serialize to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    logger.info("Welcome Pack generated successfully (%d bytes)", buffer.getbuffer().nbytes)
    return buffer.getvalue()
